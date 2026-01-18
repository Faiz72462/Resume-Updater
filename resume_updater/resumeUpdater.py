import time
from google import genai
import ast
import re
from docx import Document

import json
import datetime
import os

client = genai.Client(api_key=os.environ.get("GEMINI_API_KEY"))

def genrator(client,jd):

    response = client.models.generate_content(
        model="gemini-3-flash-preview",
        contents=f"""Prompt:

    You are a senior resume writer and ATS optimization expert.
    I will provide:

    My current Professional Summary
    "Immediately available Software Engineer based in Lucknow (Open to relocate), specialized in full stack development with Python, Django, JavaScript, HTML, CSS, REST APIs, automation, and web scraping. Bring end to end ownership of scalable, production ready applications, including real time features, subscription/payment systems, and analytics, with the ability to contribute from day one with minimal ramp up. "

    My Skills section
    "Languages: Python, JavaScript, HTML, CSS.
    Frameworks: Django, Django Channels, Bootstrap.
    APIs & Backend: REST APIs, WebSocket, Stripe API,
    Databases: PostgreSQL, SQLite, MongoDB.
    Automation & Scraping: Selenium, Beautiful Soup, Drission Page.
    Tools: Git, AWS Lambda.
    "

    My Experience section bullet points
    "SOFTWARE ENGINEER | RAFTAL TECHNOLOGIES, LUCKNOW – JUNE 2025 – DECEMBER 2025
    •	Develop and maintain full stack applications using Django, JavaScript, and modern web technologies, improving application performance and user engagement.
    •	Build cross platform desktop applications with Python (Tkinter, PyQt) for enterprise clients.
    •	Mentor junior developers through code reviews and hands on technical guidance, accelerating team productivity.
    ASSISTANT SOFTWARE ENGINEER | RAFTAL TECHNOLOGIES, LUCKNOW – JULY 2024 – JUNE 2025
    •	Designed and maintained scalable web applications, improving UI responsiveness by 35% and reducing backend load time 
    •	Automated internal workflows using Python and Selenium, reducing manual effort by 40% and increasing operational efficiency
    •	Collaborated in Agile sprints to deliver features aligned with sprint goals and business objectives.
    "
    A Job Description
    {jd}

    YOUR TASK

    Rewrite and optimize my resume content to closely match the provided Job Description, improve ATS keyword alignment, and remain 100% truthful.

    🔒 STRICT RULES (DO NOT BREAK)
    Content Rules

    ❌ Do NOT remove any existing skills, tools, technologies, or responsibilities

    ❌ Do NOT fabricate experience or exaggerate achievements

    ✅ You MAY add new skills or keywords ONLY if they are explicitly important in the Job Description AND clearly align with my existing skill set

    Added skills must be logical extensions of my background (e.g., backend, APIs, cloud, automation, scalable systems)

    Length & Structure Rules (ABSOLUTELY STRICT)

    Professional Summary

    Same number of sentences

    Same number of characters per sentence

    Skills

    Same number of skill lines

    Experience

    Same number of bullets per role

    Same number of sentences and characters per bullet

    Optimization Rules

    Use ATS-friendly keywords and terminology from the Job Description

    Prioritize technical skills, tools, frameworks, and methodologies mentioned in the JD

    No keyword stuffing — language must remain natural and professional

    OUTPUT FORMAT (MANDATORY)

    Return ONLY the following Python-style structure.
    ❌ No explanations
    ❌ No comments
    ❌ No additional text

    Output = {{
    "NEW_SUMMARY" : (
        "Sentence 1. "
        "Sentence 2. "
        "Sentence 3."
    ),

    "NEW_SKILLS" : [
        "Skill Category 1: ...",
        "Skill Category 2: ...",
        "Skill Category 3: ..."
    ],

    "EXPERIENCE_BULLETS" : {{
        "JOB TITLE 1": [
            "Bullet point 1 rewritten with strong action verbs and ATS keywords.",
            "Bullet point 2 rewritten without changing meaning or length.",
            "Bullet point 3 rewritten while preserving all skills."
        ],
        "JOB TITLE 2": [
            "Bullet point 1 rewritten.",
            "Bullet point 2 rewritten.",
            "Bullet point 3 rewritten."
        ]
    }}
    }}

    FINAL GOAL

    Produce clean, professional, production-ready resume content that:

    Scores higher in ATS

    Matches the Job Description as closely as possible

    Keeps my original experience intact

    Is immediately usable in Python scripts or automation workflows

    Do not include any commentary or explanation outside the Python output.
    """,
    )

    return response.text

def parser(response, country, company, job_portal):
    # Extract Output dictionary
    match = re.search(r'Output\s*=\s*(\{.*\})', response, re.S)
    if not match:
        raise ValueError("Output dictionary not found in response")

    parsed_output = ast.literal_eval(match.group(1))

    # Extract fields for resume update
    summary = parsed_output["NEW_SUMMARY"]
    skills = parsed_output["NEW_SKILLS"]
    experience = parsed_output["EXPERIENCE_BULLETS"]

    return summary, skills, experience, parsed_output


def resume_updator(summary, skills, experience, country, output_path):

    if country.lower() == "india":
        resume_path = "resume_templates/MOHD_FAIZ_KHAN India.docx"
    else:
        resume_path = "resume_templates/MOHD_FAIZ_KHAN Dubai.docx"

    doc = Document(resume_path)

    # ... [rest of the function logic unchanged] ...

    # -------------------------
    # NEW CONTENT (EDIT HERE)
    # -------------------------

    NEW_SUMMARY = summary

    NEW_SKILLS = skills

    EXPERIENCE_BULLETS = experience
    PROJECTS = {
        "ECommerce Platform": [
            "Built a full-stack ecommerce platform with real-time chat and secure payment integration.",
            "Implemented responsive UI and optimized checkout flow to improve conversion rates."
        ],
        "MultiMarketplace Product Aggregator": [
            "Aggregated products from multiple marketplaces with advanced filtering and search.",
            "Implemented Stripe-based subscription billing with analytics dashboards."
        ],
        "Visitor Analytics": [
            "Developed visitor tracking and clickstream analytics modules.",
            "Delivered interactive visualizations for data-driven decision-making."
        ]
    }

    # -------------------------
    # HELPER FUNCTION
    # -------------------------

    def clear_section(start_idx, end_idx):
        for i in range(start_idx, end_idx):
            doc.paragraphs[i].text = ""

    # -------------------------
    # PROCESS DOCUMENT
    # -------------------------

    paras = doc.paragraphs
    i = 0

    while i < len(paras):
        text = paras[i].text.strip()
        # print(text)

        # PROFESSIONAL SUMMARY
        if text == "PROFESSIONAL SUMMARY":
            paras[i+1].text = NEW_SUMMARY

        # SKILLS
        if text == "SKILLS":
            clear_section(i+1, i+7)
            for j, skill in enumerate(NEW_SKILLS):
                paras[i+1+j].text = skill

        # EXPERIENCE
        if "SOFTWARE ENGINEER | RAFTAL TECHNOLOGIES" in text:
            clear_section(i+1, i+4)
            try:
                for j, bullet in enumerate(EXPERIENCE_BULLETS["SOFTWARE ENGINEER | RAFTAL TECHNOLOGIES, LUCKNOW – JUNE 2025 – DECEMBER 2025"]):
                    paras[i+1+j].text = bullet
            except:
                for j, bullet in enumerate(EXPERIENCE_BULLETS["SOFTWARE ENGINEER"]):
                    paras[i+1+j].text = bullet

        if "ASSISTANT SOFTWARE ENGINEER | RAFTAL TECHNOLOGIES" in text:
            # print(text)
            clear_section(i+1, i+4)
            try:
                for j, bullet in enumerate(EXPERIENCE_BULLETS["ASSISTANT SOFTWARE ENGINEER | RAFTAL TECHNOLOGIES, LUCKNOW – JULY 2024 – JUNE 2025"]):
                    paras[i+1+j].text = bullet
            except:
                for j, bullet in enumerate(EXPERIENCE_BULLETS["ASSISTANT SOFTWARE ENGINEER"]):
                    paras[i+1+j].text = bullet

        # PROJECTS
        if "ECommerce Platform" in text:
            clear_section(i+1, i+4)
            for j, bullet in enumerate(PROJECTS["ECommerce Platform"]):
                paras[i+1+j].text = bullet

        if "MultiMarketplace Product Aggregator" in text:
            clear_section(i+1, i+4)
            for j, bullet in enumerate(PROJECTS["MultiMarketplace Product Aggregator"]):
                paras[i+1+j].text = bullet

        if "Visitor Analytics & Tracking System" in text:
            clear_section(i+1, i+4)
            for j, bullet in enumerate(PROJECTS["Visitor Analytics"]):
                paras[i+1+j].text = bullet

        i += 1

    # -------------------------
    # SAVE UPDATED RESUME
    # -------------------------

    doc.save(output_path)

def pdf_generotor(input_path, output_path):
    # PDF generation disabled due to Render limitations (requires LibreOffice)
    # convert(input_path, output_path)
    pass


